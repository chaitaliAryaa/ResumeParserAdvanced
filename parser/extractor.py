"""
extractor.py
Handles text extraction from PDF, DOCX, DOC, TXT
Falls back to PaddleOCR for scanned/image-based PDFs
"""

import io
import os
import tempfile


def extract_text(uploaded_file) -> tuple[str, str]:
    """
    Extract text from uploaded file.
    Returns: (extracted_text, method_used)
    """
    filename = uploaded_file.name.lower()
    file_bytes = uploaded_file.getvalue()

    if filename.endswith(".txt"):
        return _extract_txt(file_bytes)

    elif filename.endswith(".pdf"):
        return _extract_pdf(file_bytes)

    elif filename.endswith(".docx"):
        return _extract_docx(file_bytes)

    elif filename.endswith(".doc"):
        return _extract_doc(file_bytes)

    else:
        return "", "unsupported"


# ── TXT ──────────────────────────────────────────────────────────
def _extract_txt(file_bytes: bytes) -> tuple[str, str]:
    try:
        text = file_bytes.decode("utf-8")
    except UnicodeDecodeError:
        text = file_bytes.decode("latin-1", errors="ignore")
    return text, "plain text"


# ── PDF ──────────────────────────────────────────────────────────
def _extract_pdf(file_bytes: bytes) -> tuple[str, str]:
    # Try pdfplumber first (works for text-based PDFs)
    try:
        import pdfplumber
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            pages_text = []
            for page in pdf.pages:
                t = page.extract_text()
                if t:
                    pages_text.append(t)
            text = "\n".join(pages_text)

        if text.strip():
            return text, "pdfplumber (text PDF)"

    except Exception as e:
        print(f"pdfplumber failed: {e}")

    # Fallback: PaddleOCR for scanned PDFs
    return _extract_pdf_ocr(file_bytes)


def _extract_pdf_ocr(file_bytes: bytes) -> tuple[str, str]:
    """Convert PDF pages to images then OCR them."""
    try:
        import fitz  # PyMuPDF
        from paddleocr import PaddleOCR
        import numpy as np
        from PIL import Image

        ocr = PaddleOCR(use_angle_cls=True, lang='en', show_log=False)
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        all_text = []

        for page_num in range(len(doc)):
            page = doc[page_num]
            # Render page at 200 DPI
            mat = fitz.Matrix(200 / 72, 200 / 72)
            pix = page.get_pixmap(matrix=mat)
            img_array = np.frombuffer(pix.samples, dtype=np.uint8).reshape(
                pix.height, pix.width, pix.n
            )
            if pix.n == 4:
                img_array = img_array[:, :, :3]

            result = ocr.ocr(img_array, cls=True)
            if result and result[0]:
                lines = [word_info[1][0] for line in result for word_info in line]
                all_text.append("\n".join(lines))

        doc.close()
        return "\n".join(all_text), "PaddleOCR (scanned PDF)"

    except Exception as e:
        print(f"OCR failed: {e}")
        return "", "failed"


# ── DOCX ─────────────────────────────────────────────────────────
def _extract_docx(file_bytes: bytes) -> tuple[str, str]:
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        
        # Also extract tables
        table_texts = []
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    table_texts.append(row_text)

        full_text = "\n".join(paragraphs)
        if table_texts:
            full_text += "\n\n" + "\n".join(table_texts)

        return full_text, "python-docx"

    except Exception as e:
        print(f"DOCX extraction failed: {e}")
        return "", "failed"


# ── DOC (legacy Word) ────────────────────────────────────────────
def _extract_doc(file_bytes: bytes) -> tuple[str, str]:
    """
    For old .doc format, use antiword or textract.
    Falls back to docx2txt if available.
    """
    try:
        import subprocess
        with tempfile.NamedTemporaryFile(suffix=".doc", delete=False) as tmp:
            tmp.write(file_bytes)
            tmp_path = tmp.name

        result = subprocess.run(
            ["antiword", tmp_path],
            capture_output=True, text=True, timeout=15
        )
        os.unlink(tmp_path)

        if result.returncode == 0 and result.stdout.strip():
            return result.stdout, "antiword (.doc)"

    except Exception:
        pass

    # Fallback: try treating as docx (sometimes works)
    text, method = _extract_docx(file_bytes)
    if text.strip():
        return text, f"python-docx fallback ({method})"

    return "", "failed — install antiword for .doc support"
